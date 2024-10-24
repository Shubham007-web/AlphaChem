
import re
import os
import openai
import concurrent.futures
from docx import Document
from docx.shared import Pt
from datetime import datetime
from textstat import flesch_reading_ease


model_name = 'chatgpt-4o-latest'

api_key = "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"  # Replace with your actual API key

openai.api_key  = api_key
# Error handling for OpenAI API calls
def generate_content_with_chat(messages):
    try:
        response = openai.ChatCompletion.create(
            model=model_name,
            messages=messages,
            max_tokens=9000,
            temperature=0.7
        )
        return response['choices'][0]['message']['content'].strip()
    except openai.error.OpenAIError as e:
        print(f"OpenAI API error: {e}")
        return None

# Function to modularize and generate individual sections
def generate_lesson_section(section_prompt):
    messages = [
        {"role": "system", "content": f"You are a chemistry textbook writer for age 14-15 years."},
        {"role": "user", "content": section_prompt}
    ]
    return generate_content_with_chat(messages)


# Function to dynamically validate Flesch Reading Ease Score
def validate_reading_score(content):
    score = flesch_reading_ease(content)
    if score < 80:
        print(f"Warning: Flesch Reading Ease Score is {score}, which is lower than desired.")
    return score

# Function to clean the raw content and convert it to proper formatting in the .docx file
def clean_content_and_format(doc, raw_content):
    lines = raw_content.split("\n")
    for line in lines:
        line = line.strip()  # Remove unnecessary spaces
        
        if line.startswith("###"):
            heading = doc.add_heading(line[3:].strip(), level=3)
            run = heading.runs[0]
            run.font.size = Pt(13)
        elif line.startswith("##"):
            heading = doc.add_heading(line[2:].strip(), level=2)
            run = heading.runs[0]
            run.font.size = Pt(14)
        elif line.startswith("#"):
            heading = doc.add_heading(line[1:].strip(), level=1)
            run = heading.runs[0]
            run.font.size = Pt(16)
        elif "**" in line:
            parts = line.split("**")
            para = doc.add_paragraph()
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    para.add_run(part).font.size = Pt(12)
                else:
                    para.add_run(part).bold = True
        else:
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.font.size = Pt(12)
    
    # Validate reading score after formatting
    reading_score = validate_reading_score(raw_content)
    if reading_score < 80:
        print(f"Warning: Content might be too difficult for target audience.")

# Function to extract numbers from unit, chapter, and lesson names for file naming
def extract_number(text):
    match = re.search(r'\d+', text)
    if match:
        return match.group()
    else:
        print(f"Warning: No numbers found in the text '{text}'. Using '0' as default.")
        return '0'

# Function to add metadata to the document (e.g., generation date, version, etc.)
def add_metadata_to_doc(doc, unit_name, chapter_name, lesson_name):
    para = doc.add_paragraph()
    para.add_run(f"Generated for: {unit_name}, {chapter_name}, {lesson_name}\n").bold = True
    para.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").italic = True
    para.add_run(f"Version: 1.0\n").italic = True

# Function to create and clean the document structure, then add metadata
def create_document(unit_name, chapter_name, lesson_name, lesson_content):
    doc = Document()
    
    # Add and format Unit, Chapter, and Lesson headings
    heading1 = doc.add_heading(f'Unit: {unit_name}', level=1)
    run = heading1.runs[0]
    run.font.size = Pt(20)

    heading2 = doc.add_heading(f'Chapter: {chapter_name}', level=2)
    run = heading2.runs[0]
    run.font.size = Pt(17)

    heading3 = doc.add_heading(f'Lesson: {lesson_name}', level=3)
    run = heading3.runs[0]
    run.font.size = Pt(15)

    # Add metadata to the document
    add_metadata_to_doc(doc, unit_name, chapter_name, lesson_name)
    
    # Clean and format the generated content
    clean_content_and_format(doc, lesson_content)

    return doc

# Function to sanitize file names for saving
def sanitize_filename(filename):
    return re.sub(r'[\\/*?:"<>|]', "_", filename)

# Function to check if the folder exists and create only if it doesn't exist
def create_folder_if_not_exists(path):
    sanitized_path = sanitize_filename(path)
    if not os.path.exists(sanitized_path):
        os.makedirs(sanitized_path)

# Function to find the correct file name with versioning and sanitized naming
def find_next_version_file_name(file_path, base_file_name):
    version = 1
    base_file_name = sanitize_filename(base_file_name)
    file_name = f"{base_file_name}.docx"
    
    # Check if the file already exists, and if so, create a new version
    while os.path.exists(os.path.join(file_path, file_name)):
        version += 1
        file_name = f"V{version}_{base_file_name}.docx"
    
    return file_name

# Final function to save the generated content to a .docx file with versioning and metadata
def save_lesson_content_to_docx(unit_name, chapter_name, lesson_name, lesson_content):
    # Extract unit, chapter, and lesson numbers from their names
    unit_number = extract_number(unit_name)
    chapter_number = extract_number(chapter_name)
    lesson_number = extract_number(lesson_name)

    # Define the base folder structure
    base_folder = "AI_generated_content"
    unit_folder = f"unit{unit_number}"
    chapter_folder = f"chapter{chapter_number}"
    lesson_folder = f"lesson{lesson_number}"

    # Ensure the folder structure exists
    create_folder_if_not_exists(base_folder)
    
    # Create folder paths
    full_path = os.path.join(base_folder, unit_folder, chapter_folder, lesson_folder)
    create_folder_if_not_exists(full_path)

    # Generate the document
    doc = create_document(unit_name, chapter_name, lesson_name, lesson_content)

    # Construct the base file name in the format U1Ch1L4
    base_file_name = f'U{unit_number}Ch{chapter_number}L{lesson_number}'
    
    # Find the next available file name with versioning if necessary
    file_name = find_next_version_file_name(full_path, base_file_name)

    # Save the document with the constructed file name
    try:
        doc.save(os.path.join(full_path, file_name))
        print(f'Content successfully saved to {os.path.join(full_path, file_name)}')
    except Exception as e:
        print(f"Error saving the file {file_name}: {e}")



# Prompt Template For the various sections
def lesson_content_prompt_part1(unit_name, chapter_name, lesson_name, essential_question, lesson_vocabulary, lesson_objectives, phenomenon):
    return f"""
    Generate a detailed and structured lesson plan for "{lesson_name}" in Chapter "{chapter_name}" of Unit "{unit_name}" and please rememeber the target aduince is 13-14 years students, so don't phase "student, student have to...".
    The content should be structured, consistent, and align with the following points:
     - lesson objective: {lesson_objectives}
     - lesson vocabulary: {lesson_vocabulary}
     - Essential Question: {essential_question}

    ## Unit Title
    ## Chapter Title
    # Lesson Title
    ### Essential Questions:
     Write the following Essential Questions (EQs) to ensure they drive curiosity and critical thinking:
    - {essential_question}

    ### 1. Big Idea:
    - One line that addresses the main concept of the lesson.
    - A subordinate of the Chapter's Big Idea that addresses the main concepts in the lesson.

    ### 2. Essential Questions
    - Include the following Essential Question(s) as given:
        - {essential_question}
        - Provide answer of each questions.

    ### 3.1 Phenomenon-Based Learning
    - The lesson should build upon the chapter's storyline and introduce a specific aspect, question, or issue that will be explored through hands-on tasks. Focus on connecting the phenomenon to the lesson's tasks and investigations.
    - Phenomenon: {phenomenon}
   ### 3.2 Lesson Phenomenon
   - create lesson Phenomenon based on given unit and chapter Phenomenon and lesson onjectives.
    ### 4. Vocabulary
    - Define these key terms to support students’ understanding:
        - {lesson_vocabulary}

    ### 5. SMART Objectives
    - Write the lesson objective as given in bullet points:
        - {lesson_objectives}

    ### 6. Engage (Ignite)
    - Start with a phenomenon-related question or task to grab attention and continue on the same storyline.
    - Include one hands-on experiment relevant to the lesson topic, with a step-by-step procedure.
    - Add 2-3 follow-up questions based on the activity.
    - Provide answer of each questions if applicable

    ### 7. Pre-Explore (Direct Instruction)
    - Provide background information linking the phenomenon and key concepts.
    - Use interactive elements (notes, discussions, scaffolded questions) to break up the content.
    - Provide answer of each questions if applicable

    ### 8. Evaluate (Progress Check) - Pre-Explore
    - Frame up to 3 scaffolded questions (DOK 1-3) to connect concepts to the hands-on activity.
    - Provide answer of each questions if applicable

    ### 9. Explain (Lightbulb) 
    - Please generate approximately **5000-6000 words** of content that deeply explains the core and main concept of the lesson based on the storyline/Phenomenon: {phenomenon}.
    - Ensure the explanation follows the unit and chapter storyline and aligns with the lesson objectives : {lesson_objectives}.
    - Break down complex concepts into structured sections or subsections, making them digestible for 14-15-year-olds (around 5-6 pages long).
    - Include prompts to help students make sense of the hands-on activity by themselves, based on their prior knowledge, the evidence they gathered in their inquiry activities, and discussions with classmates.
    - After students make their attempt, provide expansion of the concepts explored in the Explore section to build a tight connection to the lessons’ objectives.
    - For every main concept explained, introduce one solved sample problem where applicable, followed by one question for students to solve as a Progress Check.
    - Provide answer of each questions if applicable
    """

def lesson_content_prompt_part2(unit_name, chapter_name, lesson_name, essential_question, lesson_vocabulary, lesson_objectives, phenomenon):
    return f"""
    ### 10. Evaluate (Progress Check) - Explain
    - Include 3 scaffolded questions (DOK 1-3) to confirm understanding of key concepts covered in the "Explain" section but don't mention  DOK level in the text,
       - Provide answer of each questions based on dok level.

    ### 11. Elaborate (Power Up)
    - Pose mini-tasks or open-ended questions encouraging deeper thinking.
    - Allow space for additional questions to extend understanding,
      - Provide answer of each questions based on dok level.

    ### 12. Final Evaluation
    - Provide 1 debate question, including arguments and points for discussion.
    - Frame 8 assessment questions:
        - 4 multiple-choice questions (with options and correct answers, and its explaination).
        - 4 long-answer questions requiring application of knowledge,
           - Provide answer of each questions.
    - Ensure alignment with the unit learning outcomes.

    ### 13. Extend (Beyond the Lesson)
    - Suggest additional tasks, readings, or challenges related to the lesson.
    - Activities (could include readings) and/or questions that challenge students to think about the application of what they’ve learned to new real-world situations, applications, or problems, enhancing their understanding of the chapter and unit tasks.
    - Provide opportunities for spaced practice, allowing students to revisit and reinforce their understanding over time.
    """


# Function to handle parallel content generation
def parallel_generate_content(unit_name, chapter_name, lesson_name, essential_question, lesson_vocabulary, lesson_objectives, phenomenon):
    # Define section prompts
    section_prompts = [
        lesson_content_prompt_part1(unit_name, chapter_name, lesson_name, essential_question, lesson_vocabulary, lesson_objectives, phenomenon),
        lesson_content_prompt_part2(unit_name, chapter_name, lesson_name, essential_question, lesson_vocabulary, lesson_objectives, phenomenon)
    ]
    
    # Parallel generation of content sections
    try:
        with concurrent.futures.ThreadPoolExecutor() as executor:
            futures = [executor.submit(generate_lesson_section, section) for section in section_prompts]
            results = [future.result() for future in concurrent.futures.as_completed(futures)]
        
        # Check if any of the results are None
        if any(result is None for result in results):
            print("Error: One of the content sections failed to generate.")
            return None
        
        # Combine results from both parts
        complete_lesson_content = "\n\n".join(results)
        return complete_lesson_content
    except Exception as e:
        print(f"Error generating lesson content: {e}")
        return None
