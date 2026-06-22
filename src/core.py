"""core.py — AlphaChemistry lesson content generator.

  * multi-provider LLM setup        -> build_llm()        (LangChain_Lesson_Generator.ipynb)
  * reading-score guidance + inputs -> module constants   (Final_Lesson_14Oct.ipynb)
  * single-pass lesson prompt       -> HUMAN_TEMPLATE      (Final_Lesson_14Oct / LangChain)
  * two-part prompt for long lessons-> PART1/PART2         (Large_Content.ipynb)
  * lesson generation               -> generate_lesson()
  * .docx export w/ folders+versions-> save_lesson_to_docx() (Unit2_Ch3_lessons.ipynb)
  * Arize Phoenix tracing           -> setup_tracing()  (auto-traces every LLM/chain call)

API keys are read from the environment (.env), never hardcoded.

Quick start:
    cp .env.example .env        # then fill in your key(s)
    python src/core.py          # generates the sample lesson into outputs/

Observability:
    Tracing is on by default (PHOENIX_TRACING=true). To view traces, run a local
    Phoenix server in another terminal:  python -m phoenix.server.main serve
    then open http://localhost:6006. To send to Phoenix Cloud instead, set
    PHOENIX_COLLECTOR_ENDPOINT (and PHOENIX_API_KEY). Set PHOENIX_TRACING=false to disable.
"""

from __future__ import annotations

import os
import re

from docx import Document
from dotenv import load_dotenv
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate

load_dotenv()

# --------------------------------------------------------------------------- #
# Configuration
# --------------------------------------------------------------------------- #
# Provider: 'ollama' (local, default) | 'azure' (Azure OpenAI) | 'openai' (standard OpenAI)
PROVIDER = os.environ.get("PROVIDER", "ollama").lower()
# Where generated lessons are written: outputs/unitX/chapterY/lessonZ/U{X}Ch{Y}L{Z}.docx
OUTPUT_ROOT = os.environ.get("OUTPUT_ROOT", "outputs")
TEMPERATURE = 0.7
MAX_TOKENS = 8192

# Readability guidance injected into the system prompt.
READING_SCORE = """Please follow these formulas in order to generate content. Target a Flesch Reading Ease Score >= 70 and grade level 8-9:
1. Flesch Reading Ease Score = 206.835 - 1.015 x (Total Words / Total Sentences) - 84.6 x (Total Syllables / Total Words)
2. Flesch-Kincaid Grade Level = 0.39 x (Total Words / Total Sentences) + 11.8 x (Total Syllables / Total Words) - 15.59
"""

# --------------------------------------------------------------------------- #
# Observability — Arize Phoenix tracing
# --------------------------------------------------------------------------- #
def _truthy(value: str | None) -> bool:
    return str(value).lower() in ("1", "true", "yes", "on")


# Cached so repeated calls (e.g. from evals.py and graph.py) register only once.
_TRACER_PROVIDER = None


def setup_tracing(project_name: str = "alphachem"):
    """Enable Arize Phoenix tracing for all LangChain LLM/chain calls.

    Auto-instruments LangChain via OpenInference, so every prompt, model call,
    token count and latency is captured with no per-call code changes. Controlled
    by env vars (see the module docstring). Safe to call even if Phoenix is not
    installed — it prints a hint and returns None instead of raising. Idempotent:
    registers once per process and returns the cached tracer provider thereafter.

    Returns the OpenTelemetry tracer provider, or None if tracing is disabled/absent.
    """
    global _TRACER_PROVIDER
    if _TRACER_PROVIDER is not None:
        return _TRACER_PROVIDER
    if not _truthy(os.environ.get("PHOENIX_TRACING", "true")):
        return None
    try:
        from phoenix.otel import register
    except ImportError:
        print(
            "[tracing] arize-phoenix not installed; skipping. Enable with:\n"
            "          pip install arize-phoenix openinference-instrumentation-langchain"
        )
        return None

    # register() reads PHOENIX_COLLECTOR_ENDPOINT / PHOENIX_API_KEY from the env;
    # auto_instrument=True activates the installed OpenInference instrumentors (LangChain).
    tracer_provider = register(project_name=project_name, auto_instrument=True)
    endpoint = os.environ.get("PHOENIX_COLLECTOR_ENDPOINT", "http://localhost:6006")
    print(f"[tracing] Phoenix tracing enabled — project={project_name!r}, view at {endpoint}")
    _TRACER_PROVIDER = tracer_provider
    return tracer_provider


# --------------------------------------------------------------------------- #
# Sample lesson inputs (example data carried over from the notebooks)
# Replace these, or load rows from data/AlphaChem_LOs.csv, to generate others.
# --------------------------------------------------------------------------- #
SAMPLE_LESSON = {
    "unit_name": "Unit 2: Atomic Structure and Bonding",
    "chapter_name": "Chapter 6: Ionic and Metallic Bonding",
    "lesson_name": "Lesson 1: Formation and Properties of Ions",
    "lesson_objective": (
        "Describe how ions are formed.\n"
        "Write the symbols and charges of ions and the octet rule and its exceptions.\n"
        "Predict the charge of an ion based on its position on the periodic table."
    ),
    "lesson_vocabulary": (
        "Octet Rule\nAnion\nCation\nElectrolyte\nElectron affinity\n"
        "Ionic radius\nIonization"
    ),
    "essential_question": (
        "How are ions formed, and what role do they play in chemical bonding?"
    ),
    "performance_expectations": (
        "HS-PS1-2: Construct and revise an explanation for the outcome of a simple "
        "chemical reaction based on the outermost electron states of atoms, trends in "
        "the periodic table, and knowledge of the patterns of chemical properties."
    ),
    "disciplinary_core_ideas": (
        "PS1.A: Structure and Properties of Matter. The structure and interactions of "
        "matter at the bulk scale are determined by electrical forces within and "
        "between atoms."
    ),
    "phenomenon": (
        "Unit phenomenon: Danger! Icy Roads. In cold northern winters, roads ice over "
        "and road salt is spread to melt the ice; the salt dissolves while metal street "
        "signs exposed to the same ice do not. Chapter phenomenon: Salt vs. Metal - why "
        "does water treat them so differently?"
    ),
}

# --------------------------------------------------------------------------- #
# LLM factory
# --------------------------------------------------------------------------- #
def build_llm(provider: str | None = None, temperature: float = TEMPERATURE,
              max_tokens: int = MAX_TOKENS):
    """Return a LangChain chat model for the chosen provider.

    'azure'  -> Azure OpenAI   (needs AZURE_OPENAI_* in .env)
    'openai' -> standard OpenAI (needs OPENAI_API_KEY)
    'ollama' -> local Ollama    (no key; model must be pulled locally)

    provider defaults to the PROVIDER env var (resolved at call time, so a
    runtime/interactive choice is honored), falling back to 'ollama'.
    """
    provider = (provider or os.environ.get("PROVIDER", "ollama")).lower()
    if provider == "azure":
        from langchain_openai import AzureChatOpenAI

        llm = AzureChatOpenAI(
            azure_endpoint=os.environ["AZURE_OPENAI_ENDPOINT"].strip(),
            azure_deployment=os.environ["AZURE_OPENAI_DEPLOYMENT"].strip(),
            api_key=os.environ["AZURE_OPENAI_API_KEY"].strip(),
            api_version=os.environ.get("OPENAI_API_VERSION", "2024-10-21").strip(),
            temperature=temperature,
            max_tokens=max_tokens,
        )
        print(f"Using Azure OpenAI deployment: {llm.deployment_name}")
        return llm

    if provider == "openai":
        from langchain_openai import ChatOpenAI

        if not os.environ.get("OPENAI_API_KEY"):
            raise RuntimeError("OPENAI_API_KEY is not set. Add it to your .env file.")
        model = os.environ.get("OPENAI_MODEL", "gpt-4o-mini")
        print(f"Using OpenAI model via LangChain: {model}")
        return ChatOpenAI(model=model, temperature=temperature, max_tokens=max_tokens)

    if provider == "ollama":
        from langchain_ollama import ChatOllama

        model = os.environ.get("OLLAMA_MODEL", "qwen2.5:14b")
        print(f"Using local Ollama model via LangChain: {model}")
        return ChatOllama(model=model, temperature=temperature, num_predict=max_tokens)

    raise ValueError(f"Unknown PROVIDER {provider!r}. Use 'azure', 'openai' or 'ollama'.")

# --------------------------------------------------------------------------- #
# Prompt templates
# --------------------------------------------------------------------------- #
SYSTEM_TEMPLATE = (
    "You are a chemistry textbook writer for ages 14-15. The content should be easy "
    "to read, so use these formulas to keep the language simple:\n{reading_score}\n"
    "Write in a way that feels human-authored. The target audience is USA school grade 9."
)

# Full 14-section lesson, generated in a single pass.
HUMAN_TEMPLATE = """Generate a detailed and structured lesson plan for "{lesson_name}" in Chapter "{chapter_name}" of Unit "{unit_name}".
The content should be structured, consistent, and align with the following points:
 - lesson objective: {lesson_objective}
 - lesson vocabulary: {lesson_vocabulary}
 - Essential Question: {essential_question}
 - Performance Expectations: {performance_expectations}
 - Disciplinary Core Ideas: {disciplinary_core_ideas}

## Unit Title
## Chapter Title
# Lesson Title

### 1. Big Idea
- One line that addresses the main concept of the lesson.
- A subordinate of the Chapter's Big Idea that addresses the main concepts in the lesson.

### 2. Essential Questions
- Include the following Essential Question(s) to encourage critical thinking:
    - {essential_question}

### 3. Phenomenon-Based Learning
- Build upon the chapter's storyline and introduce a specific aspect explored through hands-on tasks.
- Phenomenon: {phenomenon}

### 4. Vocabulary
- Define these key terms to support students' understanding:
    - {lesson_vocabulary}

### 5. SMART Objectives
- List 3-4 Specific, Measurable, Achievable, Relevant, Time-based objectives from the lesson objective:
    - {lesson_objective}

### 6. Engage (Ignite)
- Start with a phenomenon-related question or task to grab attention and continue the same story.
- Include one hands-on experiment relevant to the lesson topic, with a step-by-step procedure.
- Add 2-3 follow-up questions based on the activity.

### 7. Pre-Explore (Direct Instruction)
- Provide background information linking the phenomenon and key concepts.
- Use interactive elements (notes, discussions, scaffolded questions) to break up the content.

### 8. Evaluate (Progress Check) - Pre-Explore
- Frame up to 3 scaffolded questions (DOK 1-3) to connect concepts to the hands-on activity.

### 9. Explore (Pathfinder)
- Guide students through a hands-on activity with clear instructions.
- Ensure they collect data and engage in group discussions.
- Use retrieval practice (quizzes or questions) to reinforce learning.

### 10. Explain (Lightbulb)
- This section should be around 4500 words and explain the core concept of the lesson based on the storyline.
- Align with the lesson objective and the Big Idea; break complex concepts into digestible sections for 14-15 year olds.
- For every main concept, include one sample solved problem where applicable, then one question for students to solve as a Progress Check.

### 11. Evaluate (Progress Check) - Explain
- Include 3 scaffolded questions (DOK 1-3) to confirm understanding of key concepts covered in the Explain section.

### 12. Elaborate (Power Up)
- Pose mini-tasks or open-ended questions encouraging deeper thinking.

### 13. Final Evaluation
- Provide 1 debate question, including arguments and points for discussion.
- Frame 8 assessment questions:
    - 4 multiple-choice questions (with options and correct answers).
    - 4 long-answer questions requiring application of knowledge.
- Ensure alignment with the unit learning outcomes.

### 14. Extend (Beyond the Lesson) [Optional]
- Suggest additional tasks, readings, or challenges related to the lesson.
- Provide opportunities for spaced practice to reinforce previously learned key concepts.
"""

# Two-part split for very long lessons (the "Large Content" approach): the heavy
# "Explain" section is generated separately so neither call hits the token ceiling.
HUMAN_TEMPLATE_PART1 = """Generate the first half of a detailed lesson plan for "{lesson_name}" in Chapter "{chapter_name}" of Unit "{unit_name}".
Align with:
 - lesson objective: {lesson_objective}
 - lesson vocabulary: {lesson_vocabulary}
 - Essential Question: {essential_question}
 - Phenomenon: {phenomenon}

Produce sections 1-10:
1. Big Idea
2. Essential Questions (answer each)
3. Phenomenon-Based Learning
4. Vocabulary
5. SMART Objectives
6. Engage (Ignite) - include a hands-on experiment and 2-3 follow-up questions
7. Pre-Explore (Direct Instruction)
8. Evaluate (Progress Check) - Pre-Explore - 3 scaffolded questions (DOK 1-3)
9. Explore (Pathfinder) - hands-on activity with data collection and retrieval practice
10. Explain (Lightbulb) - ~4500 words of core content aligned to the storyline, with a solved example and Progress Check per main concept.
"""

HUMAN_TEMPLATE_PART2 = """Continue the lesson plan for "{lesson_name}" in Chapter "{chapter_name}" of Unit "{unit_name}".
Phenomenon: {phenomenon}

Produce sections 11-14:
11. Evaluate (Progress Check) - Explain - 3 scaffolded questions (DOK 1-3) with answers.
12. Elaborate (Power Up) - mini-tasks / open-ended questions with answers.
13. Final Evaluation - 1 debate question plus 8 assessment questions
    (4 MCQs with options, correct answers and explanations; 4 long-answer with answers),
    aligned with the unit learning outcomes.
14. Extend (Beyond the Lesson) - additional tasks/readings and spaced-practice opportunities.
"""


def _chain(human_template: str, llm):
    """Build a system+human -> llm -> string chain for a given human template."""
    prompt = ChatPromptTemplate.from_messages(
        [("system", SYSTEM_TEMPLATE), ("human", human_template)]
    )
    return prompt | llm | StrOutputParser()


# --------------------------------------------------------------------------- #
# Generation
# --------------------------------------------------------------------------- #
def generate_lesson(lesson: dict, llm=None, two_part: bool = False,
                    provider: str | None = None) -> str:
    """Generate lesson content via the LangGraph workflow and return the text.

    Backed by graph.py (planner -> part1 -> part2 -> part3 -> summary), which
    replaced the old single-chain implementation. Returns the assembled lesson
    without evaluating or exporting it (use graph.run_lesson_graph() for the full
    pipeline). The ``llm`` and ``two_part`` arguments are accepted for backward
    compatibility; the multi-node graph supersedes the old two-part split.
    """
    from graph import generate_content  # lazy import avoids a circular import

    return generate_content(lesson, provider=provider)


# --------------------------------------------------------------------------- #
# File management: Unit/Chapter/Lesson folders + automatic versioning
# --------------------------------------------------------------------------- #
def extract_number(text: str) -> str:
    """Pull the first integer out of a name, e.g. 'Unit 2: ...' -> '2'."""
    match = re.search(r"\d+", text)
    if match:
        return match.group()
    print(f"Warning: no number found in {text!r}; using '0'.")
    return "0"


def find_next_version_file_name(folder: str, base_file_name: str) -> str:
    """Return base_file_name.docx, or V2_/V3_... if earlier versions already exist."""
    version = 1
    file_name = f"{base_file_name}.docx"
    while os.path.exists(os.path.join(folder, file_name)):
        version += 1
        file_name = f"V{version}_{base_file_name}.docx"
    return file_name


def create_document(unit_name: str, chapter_name: str, lesson_name: str,
                    lesson_content: str) -> Document:
    """Build a .docx with unit/chapter/lesson headings and the generated body."""
    doc = Document()
    doc.add_heading(f"Unit: {unit_name}", level=1)
    doc.add_heading(f"Chapter: {chapter_name}", level=2)
    doc.add_heading(f"Lesson: {lesson_name}", level=3)
    doc.add_paragraph(lesson_content)
    return doc


def save_lesson_to_docx(unit_name: str, chapter_name: str, lesson_name: str,
                        lesson_content: str, output_root: str = OUTPUT_ROOT) -> str:
    """Save content to outputs/unitX/chapterY/lessonZ/U{X}Ch{Y}L{Z}.docx (versioned).

    Returns the full path written.
    """
    unit_no = extract_number(unit_name)
    chapter_no = extract_number(chapter_name)
    lesson_no = extract_number(lesson_name)

    folder = os.path.join(
        output_root, f"unit{unit_no}", f"chapter{chapter_no}", f"lesson{lesson_no}"
    )
    os.makedirs(folder, exist_ok=True)

    base_file_name = f"U{unit_no}Ch{chapter_no}L{lesson_no}"
    file_name = find_next_version_file_name(folder, base_file_name)
    full_path = os.path.join(folder, file_name)

    create_document(unit_name, chapter_name, lesson_name, lesson_content).save(full_path)
    print(f"Content saved to {full_path}")
    return full_path


# --------------------------------------------------------------------------- #
# CLI entry point
# --------------------------------------------------------------------------- #
def main(lesson: dict = SAMPLE_LESSON, two_part: bool = False,
         provider: str | None = None, run_eval: bool = True) -> str:
    """Run the full LangGraph pipeline end-to-end and return the saved .docx path.

    planner -> part1 -> part2 -> part3 -> summary -> evaluation -> (revision) ->
    export. The export node writes the versioned .docx; evaluation (when enabled)
    scores the lesson and logs to Phoenix. ``two_part`` is accepted for backward
    compatibility but no longer changes behavior (the graph always splits work).
    """
    from graph import run_lesson_graph  # lazy import avoids a circular import

    state = run_lesson_graph(lesson, provider=provider, run_eval=run_eval)
    print(state.get("final_lesson_content", ""))
    results = state.get("evaluation_results") or {}
    if results:
        import evals

        evals.print_report(results)
    return state.get("export_path", "")


def _ask(prompt_text: str, options: dict, default_key: str):
    """Show a small numbered menu and return the chosen value."""
    print(prompt_text)
    for key, (label, _value) in options.items():
        marker = "  (default)" if key == default_key else ""
        print(f"  {key}) {label}{marker}")
    choice = input(f"Choice [default {default_key}]: ").strip() or default_key
    label, value = options.get(choice, options[default_key])
    print(f"-> {label}\n")
    return value


def _interactive_menu():
    """Prompt for provider and generation mode at runtime."""
    provider = _ask(
        "Select LLM provider:",
        {
            "1": ("ollama - local model", "ollama"),
            "2": ("azure  - Azure OpenAI", "azure"),
            "3": ("openai - OpenAI", "openai"),
        },
        default_key="1",
    )
    two_part = _ask(
        "Select generation mode:",
        {
            "1": ("single pass (HUMAN_TEMPLATE)", False),
            "2": ("two iterations (HUMAN_TEMPLATE_PART1 + PART2)", True),
        },
        default_key="1",
    )
    return provider, two_part


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(
        description="Generate a chemistry lesson and save it as a .docx."
    )
    parser.add_argument("--provider", choices=["azure", "openai", "ollama"],
                        help="Skip the menu and use this provider.")
    parser.add_argument("--two-part", action="store_true",
                        help="Accepted for backward compatibility (no longer changes behavior).")
    parser.add_argument("--no-eval", action="store_true",
                        help="Skip the evaluation + revision stage (generate and export only).")
    args = parser.parse_args()

    # Use flags if given; otherwise ask interactively.
    if args.provider or args.two_part or args.no_eval:
        chosen_provider = args.provider or os.environ.get("PROVIDER", "ollama")
        chosen_two_part = args.two_part
    else:
        chosen_provider, chosen_two_part = _interactive_menu()

    main(provider=chosen_provider, two_part=chosen_two_part, run_eval=not args.no_eval)
